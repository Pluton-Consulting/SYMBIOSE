"""
Banc « reproduire un document à l'identique » — 02/09.

Demande de Noa : « concernant les docs Excel, Word, etc., il doit être capable
en analysant des docs de les reproduire à l'identique, soit en copiant soit en
téléchargeant une copie et en remplaçant le contenu ».

CE BANC FABRIQUE DE VRAIS FICHIERS, LES REMPLIT, PUIS LES ROUVRE POUR VÉRIFIER
QUE LA MISE EN FORME A SURVÉCU. C'est la seule preuve qui vaille ici : un
contrôle sur le source dirait que le code appelle `doc.save()`, il ne dirait
pas que le logo est encore là. On construit donc un document Word avec ce qui
casse en pratique — en-tête, pied de page, tableau, gras, couleur, texte
éclaté en plusieurs runs — et un classeur Excel avec des formules et des
largeurs de colonnes réglées.

LE DÉFAUT QU'IL EXISTE POUR ATTRAPER : dans un .docx, Word découpe un
paragraphe en « runs » à chaque changement de mise en forme, et en crée aussi
pour ses propres raisons (correction, copier-coller, reprise de frappe). « Devis
n° DEV-2025-014 » vit donc souvent en plusieurs runs, et chercher la référence
dans chacun ne trouve RIEN. C'est exactement sur les documents retouchés à la
main — les vrais — qu'un remplacement naïf échoue en silence. Le contrôle 5 le
reproduit exprès.

Prérequis : python-docx et openpyxl. Ils sont dans requirements.txt (le
conteneur les a) mais pas forcément sur la machine de développement : le banc
le dit et s'arrête proprement plutôt que de tomber.
"""
import io
import pathlib
import sys

BACKEND = pathlib.Path(sys.argv[1] if len(sys.argv) > 1 else "backend")
echecs = []


def verifier(nom, cond, detail=""):
    print(f"  {'✓' if cond else '✗'} {nom}" + (f"  → {detail}" if detail and not cond else ""))
    if not cond:
        echecs.append(nom)


print(f"\n═══ REPRODUIRE UN DOCUMENT — {BACKEND.resolve().parent}\n")

try:
    import docx
    from docx.shared import Pt, RGBColor
    from openpyxl import Workbook, load_workbook
except ImportError as e:
    print(f"  ⚠ {e.name} absent de cette machine : banc non joué.")
    print("    (il est dans requirements.txt ; le conteneur l'a)\n")
    sys.exit(0)

import importlib.util  # noqa: E402

spec = importlib.util.spec_from_file_location(
    "trame_banc", BACKEND / "bureautique" / "trame.py")
tr = importlib.util.module_from_spec(spec)
spec.loader.exec_module(tr)


# ── Un document Word avec tout ce qui casse en pratique ──────────────────
def fabriquer_docx() -> bytes:
    d = docx.Document()
    section = d.sections[0]

    # L'EN-TÊTE porte l'identité du document : c'est ce qu'on oublie le plus
    # souvent, et ce dont l'absence se voit le plus.
    entete = section.header.paragraphs[0]
    entete.text = "SYMBIOSE PAYSAGE — Devis DEV-2025-014"
    section.footer.paragraphs[0].text = "TVA FR12345678901 — page"

    titre = d.add_paragraph()
    r = titre.add_run("Devis pour Monsieur Dupont")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    # UN TEXTE ÉCLATÉ EN CINQ RUNS : le cas qui fait échouer un remplacement
    # naïf. Aucun run ne contient « DEV-2025-014 » en entier.
    p = d.add_paragraph()
    for fragment in ("Référence ", "DEV", "-2025", "-", "014 du 3 avril 2026"):
        p.add_run(fragment)

    t = d.add_table(rows=2, cols=2)
    t.style = "Light Grid Accent 1"
    t.cell(0, 0).text = "Poste"
    t.cell(0, 1).text = "Montant"
    t.cell(1, 0).text = "Terrasse bois pour Monsieur Dupont"
    t.cell(1, 1).text = "4 200 EUR"

    sortie = io.BytesIO()
    d.save(sortie)
    return sortie.getvalue()


def fabriquer_xlsx() -> bytes:
    w = Workbook()
    f = w.active
    f.title = "Devis"
    f["A1"] = "Client"
    f["B1"] = "Monsieur Dupont"
    f["A2"] = "Référence"
    f["B2"] = "DEV-2025-014"
    f["A3"] = "Quantité"
    f["B3"] = 12
    f["A4"] = "Prix unitaire"
    f["B4"] = 350
    f["A5"] = "Total"
    f["B5"] = "=B3*B4"          # une FORMULE, qui doit survivre
    f.column_dimensions["A"].width = 28.5   # une largeur réglée à la main
    f["A1"].font = f["A1"].font.copy(bold=True)
    sortie = io.BytesIO()
    w.save(sortie)
    return sortie.getvalue()


ORIGINAL_DOCX = fabriquer_docx()
ORIGINAL_XLSX = fabriquer_xlsx()

# ── 1. RECONNAÎTRE UNE TRAME ─────────────────────────────────────────────
verifier("un .docx est reconnu", tr.type_de("devis.docx") == "docx")
verifier("un .xlsx aussi", tr.type_de("suivi.xlsx") == "xlsx")
verifier("un .xlsm aussi (classeur à macros)", tr.type_de("suivi.xlsm") == "xlsx")
verifier("le type MIME suffit quand le nom ne dit rien",
         tr.type_de("piece", tr.TYPES["docx"]) == "docx")
# UN PDF N'EST PAS UNE TRAME REMPLISSABLE, et c'est volontaire : on ne modifie
# pas un PDF sans le reconstruire, donc sans perdre ce qu'on voulait garder.
verifier("un PDF n'est PAS une trame remplissable", tr.type_de("devis.pdf") is None)
verifier("ni une image", tr.type_de("logo.png") is None)

# ── 2. ANALYSER : ce que la trame contient ───────────────────────────────
a = tr.analyser(ORIGINAL_DOCX, "docx")
verifier("EXÉCUTÉ — l'analyse compte les tableaux", a["tableaux"] == 1, str(a.get("tableaux")))
verifier("elle voit l'en-tête (c'est lui qui porte l'identité)", a["entete"] is True)
verifier("elle lit le texte de l'en-tête, pas seulement du corps",
         any("SYMBIOSE PAYSAGE" in t for t in a["textes"]))
verifier("et le contenu des tableaux",
         any("Terrasse bois" in t for t in a["textes"]))
ok, pourquoi = tr.exploitable(a)
verifier("un document qui porte du texte est exploitable", ok, pourquoi)

vide = docx.Document()
tampon = io.BytesIO()
vide.save(tampon)
ok, pourquoi = tr.exploitable(tr.analyser(tampon.getvalue(), "docx"))
verifier("un document SANS texte est refusé, avec sa raison",
         ok is False and "scan" in pourquoi)

# LES VARIABLES SE RECONNAISSENT, ELLES NE S'INVENTENT PAS. On ne devine
# jamais qu'un mot est une variable parce qu'il ressemble à un nom : remplacer
# « Dupont » partout détruirait « rue Dupont » et « société Dupont & Fils ».
d2 = docx.Document()
d2.add_paragraph("Devis pour {client}, chantier [[ville]], le {date}.")
tampon = io.BytesIO()
d2.save(tampon)
v = tr.analyser(tampon.getvalue(), "docx")["variables"]
verifier("les variables {…} et [[…]] sont reconnues",
         v == ["client", "ville", "date"], str(v))
verifier("un document sans variable n'en invente pas",
         tr.analyser(ORIGINAL_DOCX, "docx")["variables"] == [])

x = tr.analyser(ORIGINAL_XLSX, "xlsx")
verifier("un classeur rend ses feuilles et leurs dimensions",
         x["feuilles"][0]["nom"] == "Devis" and x["feuilles"][0]["lignes"] == 5,
         str(x.get("feuilles")))

# ── 3. REMPLIR UN WORD, ET VÉRIFIER CE QUI A SURVÉCU ─────────────────────
TABLE = {"Monsieur Dupont": "Madame Martin",
         "DEV-2025-014": "DEV-2026-088",
         "4 200 EUR": "5 750 EUR"}
rendu, faits = tr.remplir(ORIGINAL_DOCX, "docx", TABLE)
verifier("EXÉCUTÉ — le document est réécrit", faits > 0, str(faits))

relu = docx.Document(io.BytesIO(rendu))
corps = "\n".join(p.text for p in relu.paragraphs)
tableau = "\n".join(c.text for t in relu.tables for l in t.rows for c in l.cells)
entete = relu.sections[0].header.paragraphs[0].text
pied = relu.sections[0].footer.paragraphs[0].text

verifier("le nom du client est remplacé dans le corps", "Madame Martin" in corps)
verifier("et dans le TABLEAU", "Madame Martin" in tableau)
verifier("l'ancien nom a bien disparu",
         "Dupont" not in corps and "Dupont" not in tableau and "Dupont" not in entete)
verifier("le montant du tableau est remplacé", "5 750 EUR" in tableau)
# CE CONTRÔLE EST LE PLUS IMPORTANT DE LA SECTION : l'en-tête et le pied de
# page portent la référence et le numéro de TVA. Les oublier produirait un
# document qui se contredit, l'ancienne référence subsistant en haut de chaque
# page — le genre d'erreur qui part chez un client.
verifier("L'EN-TÊTE est remplacé lui aussi", "DEV-2026-088" in entete, entete)
verifier("le pied de page est intact (rien ne le visait)", "FR12345678901" in pied)

# ── 4. LA MISE EN FORME A SURVÉCU ────────────────────────────────────────
titre = relu.paragraphs[0]
verifier("le gras du titre a survécu", titre.runs[0].bold is True)
verifier("sa taille aussi", titre.runs[0].font.size == Pt(18))
verifier("sa couleur aussi", str(titre.runs[0].font.color.rgb) == "1B5E20")
verifier("le style du tableau a survécu",
         relu.tables[0].style.name == "Light Grid Accent 1",
         relu.tables[0].style.name)
verifier("la structure est identique (mêmes compteurs)",
         len(relu.tables) == len(docx.Document(io.BytesIO(ORIGINAL_DOCX)).tables))

# ── 5. LE PIÈGE DES RUNS ─────────────────────────────────────────────────
# « DEV-2025-014 » était éclaté en quatre runs dans le document d'origine :
# aucun ne le contenait en entier. Un remplacement run par run — l'implémentation
# naïve, celle qu'on écrit d'abord — n'aurait RIEN trouvé et n'aurait rien dit.
eclate = [p for p in relu.paragraphs if p.text.startswith("Référence")]
verifier("LE PIÈGE DES RUNS — un texte éclaté sur plusieurs runs est remplacé",
         eclate and "DEV-2026-088" in eclate[0].text,
         eclate[0].text if eclate else "paragraphe introuvable")
verifier("et la date qui suivait dans le même paragraphe est conservée",
         eclate and "3 avril 2026" in eclate[0].text)

# ── 6. REMPLIR UN EXCEL, FORMULES COMPRISES ──────────────────────────────
rendu_x, faits_x = tr.remplir(ORIGINAL_XLSX, "xlsx", TABLE)
verifier("EXÉCUTÉ — le classeur est réécrit", faits_x > 0, str(faits_x))
relu_x = load_workbook(io.BytesIO(rendu_x), data_only=False)
f = relu_x["Devis"]
verifier("la valeur texte est remplacée", f["B1"].value == "Madame Martin")
verifier("la référence aussi", f["B2"].value == "DEV-2026-088")
# LA FORMULE EST LE CONTRÔLE QUI COMPTE : `load_workbook(data_only=True)`
# rendrait la dernière valeur calculée par Excel et l'écrirait en dur. Le
# classeur serait mort, ses totaux figés, et personne ne le verrait avant
# d'avoir modifié une quantité.
verifier("LA FORMULE A SURVÉCU (data_only=False)", f["B5"].value == "=B3*B4",
         str(f["B5"].value))
verifier("les nombres n'ont pas été convertis en texte", f["B3"].value == 12)
verifier("la largeur de colonne réglée à la main a survécu",
         abs((f.column_dimensions["A"].width or 0) - 28.5) < 0.01,
         str(f.column_dimensions["A"].width))
verifier("le gras d'une cellule a survécu", f["A1"].font.bold is True)
verifier("le nom de la feuille est conservé", relu_x.sheetnames == ["Devis"])

# ── 7. LES REFUS ─────────────────────────────────────────────────────────
for mauvaise, pourquoi in (({}, "table vide"),
                           ({"": "x"}, "clé vide : remplacerait partout"),
                           ({"  ": "x"}, "clé blanche"),
                           (None, "pas une table")):
    try:
        tr.remplir(ORIGINAL_DOCX, "docx", mauvaise)
        refuse = False
    except (ValueError, TypeError, AttributeError):
        refuse = True
    verifier(f"refusé : {pourquoi}", refuse)

try:
    tr.remplir(ORIGINAL_DOCX, "docx", {f"k{i}": "v" for i in range(tr.MAX_REMPLACEMENTS + 1)})
    borne = False
except ValueError:
    borne = True
verifier("au-delà du plafond de remplacements, refus explicite", borne)

try:
    tr.remplir(ORIGINAL_DOCX, "pdf", TABLE)
    refuse = False
except ValueError:
    refuse = True
verifier("un type non géré est refusé, pas traité au hasard", refuse)

# ── 8. RIEN À REMPLACER N'EST PAS UNE ERREUR ─────────────────────────────
# Un document où le texte cherché n'apparaît pas doit sortir INTACT, pas
# échouer : c'est le cas d'une trame qu'on remplit partiellement.
rendu2, faits2 = tr.remplir(ORIGINAL_DOCX, "docx", {"introuvable": "x"})
verifier("un texte absent ne fait rien, et ne lève pas", faits2 == 0)
verifier("et le document rendu reste lisible",
         "Dupont" in "\n".join(p.text for p in docx.Document(io.BytesIO(rendu2)).paragraphs))

# ── 9. LES SKILLS, exécutés contre une base doublée ──────────────────────
import json as _json  # noqa: E402
import types as _types  # noqa: E402

TRAMES = [
    {"id": "1", "nom": "Devis terrasse", "genre": "document", "type_fichier": "docx",
     "contenu": ORIGINAL_DOCX, "texte": "", "description": "devis bois",
     "variables": "[]", "actif": True},
    {"id": "2", "nom": "Devis élagage", "genre": "document", "type_fichier": "docx",
     "contenu": ORIGINAL_DOCX, "texte": "", "description": "devis arbres",
     "variables": "[]", "actif": True},
    {"id": "3", "nom": "Montage appel offres", "genre": "methode", "type_fichier": None,
     "contenu": None, "texte": "1. Lire le RC. 2. Vérifier les délais.",
     "description": "process AO", "variables": "[]", "actif": True},
]
ECRITS = []


class _Conn:
    async def fetch(self, sql, *a):
        if "lower(nom) = lower($1)" in sql:
            return [t for t in TRAMES if t["nom"].lower() == (a[0] or "").lower()]
        if "nom ILIKE" in sql:
            m = (a[0] or "").strip("%").lower()
            return [t for t in TRAMES if m in t["nom"].lower()]
        return [dict(t, usages=0, octets=0, derniere_maj=None) for t in TRAMES]

    async def fetchval(self, sql, *a):
        return len(TRAMES)

    async def execute(self, sql, *a):
        ECRITS.append(" ".join(sql.split()))


class _Db:
    async def __aenter__(self):
        return _Conn()

    async def __aexit__(self, *x):
        return False


_faux = _types.ModuleType("database.connection")
_faux.get_db = lambda: _Db()
_paquet = _types.ModuleType("database")
_paquet.__path__ = []
sys.modules.setdefault("database", _paquet)
sys.modules["database.connection"] = _faux
sys.path.insert(0, str(BACKEND))

import asyncio  # noqa: E402

spec_s = importlib.util.spec_from_file_location(
    "trames_skills", BACKEND / "skills" / "trames.py")
sk = importlib.util.module_from_spec(spec_s)
try:
    spec_s.loader.exec_module(sk)
    charge = True
except Exception as e:  # noqa: BLE001
    charge = False
    print(f"  ⚠ skills/trames.py non chargeable ici ({type(e).__name__}: {e})")

if charge:
    class _Moi:
        id = "u-1"
        email = "noa@example.fr"

    r = asyncio.run(sk.mes_trames({}, _Moi()))
    verifier("EXÉCUTÉ — les trames se listent", r["nombre"] == 3, str(r.get("nombre")))
    verifier("le tableau s'affiche MÉCANIQUEMENT (leçon de fd1bcf7)",
             r.get("bloc_garanti") and r["bloc_ui"]["type"] == "table")
    verifier("et le modèle a INTERDICTION de le recopier",
             "n'écris aucun bloc" in r["a_faire"])

    # ON NE DEVINE JAMAIS LAQUELLE : « devis » vise deux trames.
    r = asyncio.run(sk.utiliser_trame({"trame": "devis",
                                       "remplacements": {"a": "b"}}, _Moi()))
    verifier("LE GARDE-FOU — « devis » vise deux trames : AUCUNE n'est reprise",
             r.get("ambigu") is True and len(r["candidates"]) == 2)
    verifier("et le modèle ne tranche pas à la place de la personne",
             "Ne choisis PAS à sa place" in r["a_faire"])

    # Une trame NOMMÉE se reprend, et le document sort rempli.
    depose = {}
    _atelier = _types.ModuleType("bureautique.atelier")
    _atelier.deposer_fichier = lambda nom, octets, prop, origine="depot": (
        depose.update({"nom": nom, "octets": octets}) or "jeton-x")
    sys.modules["bureautique.atelier"] = _atelier
    # `from bureautique import atelier` lit l'ATTRIBUT du paquet, pas
    # sys.modules : sans cette ligne, le vrai atelier écrivait sur le disque et
    # la doublure ne voyait rien passer.
    import bureautique as _paquet_bureau  # noqa: E402
    _paquet_bureau.atelier = _atelier
    r = asyncio.run(sk.utiliser_trame(
        {"trame": "Devis terrasse",
         "remplacements": {"Monsieur Dupont": "Madame Martin"}}, _Moi()))
    verifier("EXÉCUTÉ — une trame nommée est reprise", r["remplacements"] > 0, str(r))
    verifier("le fichier produit s'affiche MÉCANIQUEMENT",
             r.get("bloc_garanti") and r["bloc_ui"]["type"] == "fichier")
    verifier("et le document depose porte vraiment le remplacement",
             "Madame Martin" in "\n".join(
                 p.text for p in docx.Document(io.BytesIO(depose["octets"])).paragraphs))
    verifier("l'usage est compté (une trame jamais reprise est à retirer)",
             any("usages = usages + 1" in e for e in ECRITS))

    # RIEN À REMPLACER SE DIT. Rendre un document identique sans le signaler
    # ferait croire au travail fait — le genre de silence qui part chez un client.
    r = asyncio.run(sk.utiliser_trame({"trame": "Devis terrasse",
                                       "remplacements": {"absent": "x"}}, _Moi()))
    verifier("aucun texte trouvé : le document sort quand même, et on le DIT",
             r["remplacements"] == 0 and "Aucun des textes cherchés" in r["message_final"])

    # Une méthode se lit, elle ne se remplit pas.
    r = asyncio.run(sk.utiliser_trame({"trame": "Montage appel offres"}, _Moi()))
    verifier("une méthode rend son texte, sans remplacement",
             r["genre"] == "methode" and "Lire le RC" in r["texte"])
    verifier("et le modèle doit la SUIVRE, pas la recopier", "suis-la" in r["a_faire"])

    # Les refus.
    for params, pourquoi in (
            ({}, "sans nom"),
            ({"nom": "x", "genre": "inconnu"}, "genre hors des trois"),
            ({"nom": "x", "genre": "methode"}, "méthode sans texte"),
            ({"nom": "x", "genre": "document"}, "document sans fichier")):
        try:
            asyncio.run(sk.enregistrer_trame(params, _Moi()))
            refuse = False
        except sk.TrameInvalide:
            refuse = True
        verifier(f"refusé : {pourquoi}", refuse)

    verifier("les quatre gestes sont déclarés EN UN SEUL endroit (le registre)",
             set(sk.SKILLS) == {"enregistrer_trame", "mes_trames",
                                "utiliser_trame", "oublier_trame"},
             str(sorted(sk.SKILLS)))
    verifier("lister est une LECTURE, le reste écrit dans l'app seulement",
             sk.SKILLS["mes_trames"].effet == "lecture"
             and all(sk.SKILLS[n].effet == "ecriture_interne"
                     for n in ("enregistrer_trame", "utiliser_trame", "oublier_trame")))
    verifier("chacun porte son libellé « je … » pour l'écran",
             all(d.libelle.startswith("je ") for d in sk.SKILLS.values()))
    # LE MODÈLE NE DÉSIGNE QUE CE QU'UN GESTE LUI A RENDU : la règle de
    # `mail/attaches.py`, redite au catalogue pour qu'il ne compose pas un chemin.
    verifier("le catalogue interdit de composer un chemin de fichier",
             "jamais un chemin que tu composes"
             in sk.SKILLS["enregistrer_trame"].description)

# La migration qui porte tout cela.
mig = BACKEND / "database" / "migrations" / "033_trames.sql"
verifier("la migration 033 existe", mig.exists())
if mig.exists():
    sql = mig.read_text(encoding="utf-8")
    verifier("elle est idempotente (règle du projet)",
             "CREATE TABLE IF NOT EXISTS trames" in sql)
    verifier("le nom d'une trame est unique, insensible à la casse",
             "ON trames (lower(nom))" in sql)
    verifier("les octets vivent en base, pas sur un volume purgé à 24 h",
             "contenu       BYTEA" in sql and "purgé à 24 h" in sql)

print(f"\n{'═' * 70}\n{'✗ ' + str(len(echecs)) + ' échec(s) : ' + ', '.join(echecs) if echecs else '✓ 0 échec'}\n")
sys.exit(1 if echecs else 0)
