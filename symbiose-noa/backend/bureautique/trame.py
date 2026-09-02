"""
REPRODUIRE UN DOCUMENT À L'IDENTIQUE — Word et Excel.

DEMANDE DE NOA (02/09) : « concernant les docs Excel, Word, etc., il doit être
capable en analysant des docs de les reproduire à l'identique, soit en copiant
soit en téléchargeant une copie et en remplaçant le contenu ; et il doit être
capable d'enregistrer des trames qu'il reprend à chaque fois, que ce soit pour
des documents, logo, méthodes, process. »

POURQUOI CE MODULE NE RESSEMBLE PAS À `rendu.py`. Le rendu existant part d'un
document VIERGE (`Document()`, `Workbook()`) et repose un contenu à partir d'un
vocabulaire de blocs. C'est le bon outil pour fabriquer un document neuf, et
c'est le mauvais pour en reproduire un : il ne sait poser que ce que le
vocabulaire nomme, et le devis d'un client porte cent choses qu'aucun
vocabulaire ne nommera jamais — un logo à sa place exacte, une trame de
tableau, une police maison, un pied de page avec un numéro de TVA, des largeurs
de colonnes réglées à la main.

D'OÙ LE PRINCIPE, ET IL TIENT EN UNE PHRASE : ON N'ÉCRIT PAS LE DOCUMENT, ON
OUVRE L'ORIGINAL ET ON N'EN CHANGE QUE LE TEXTE. Tout ce qu'on ne touche pas
reste par construction — styles, en-têtes, images, mise en page, formules. Ce
n'est pas « presque à l'identique », c'est le fichier lui-même.

LE PIÈGE DE WORD, ET C'EST LE CŒUR DU MODULE. Dans un .docx, un paragraphe est
découpé en « runs », un par changement de mise en forme. Word en crée aussi
pour des raisons qui lui appartiennent : une correction orthographique, un
copier-coller, un retour de frappe. « Devis n° DEV-2025-014 » peut donc vivre
en cinq runs, et chercher « DEV-2025-014 » dans chacun d'eux ne trouve RIEN. Un
remplacement naïf échoue silencieusement sur les documents réels, précisément
ceux qui ont été retouchés à la main. On travaille donc sur le texte ENTIER du
paragraphe, puis on repose le résultat dans le premier run et on vide les
autres : la mise en forme du début du paragraphe l'emporte, ce qui est le
comportement attendu quand on remplace une valeur dans une phrase.

CE QU'ON NE FAIT PAS, ET POURQUOI. On n'exécute jamais de code produit par un
modèle — même décision que `modele.py` : le modèle fournit une TABLE de
remplacements (texte → texte), rien d'autre. Il ne choisit ni les styles, ni la
structure, ni les fichiers ouverts.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Optional

logger = logging.getLogger("symbiose.trame")

# Un document dont on n'extrait aucun texte n'est pas une trame exploitable :
# c'est un scan, un PDF déguisé, ou un fichier qu'on ne sait pas lire. On le
# dit plutôt que de l'enregistrer et de le voir échouer au premier usage.
MIN_TEXTE_UTILE = 20

# Bornes de sûreté. Une trame vit en base : elle doit rester une trame, pas un
# dépôt de fichiers. 12 Mo couvre très largement un devis avec logo et photos.
MAX_OCTETS = 12 * 1024 * 1024
MAX_REMPLACEMENTS = 200

# Ce qu'on sait rouvrir et réécrire. Le PDF n'y est pas, et c'est volontaire :
# on ne modifie pas un PDF sans le reconstruire, donc sans perdre ce qu'on
# cherchait à garder. Un PDF se garde comme pièce, pas comme trame à remplir.
TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def type_de(nom: str, mime: str = "") -> Optional[str]:
    """« docx », « xlsx », ou None si ce n'est pas une trame remplissable."""
    n = (nom or "").lower().strip()
    if n.endswith(".docx") or "wordprocessingml" in (mime or ""):
        return "docx"
    if n.endswith((".xlsx", ".xlsm")) or "spreadsheetml" in (mime or ""):
        return "xlsx"
    return None


# ── Analyser : ce que le document contient, et ce qui s'y remplace ───────

# UNE VARIABLE SE RECONNAÎT, ELLE NE S'INVENTE PAS. Trois écritures couvrent ce
# qu'on rencontre en pratique : {client}, [[client]] et «client». On ne devine
# JAMAIS qu'un mot est une variable parce qu'il ressemble à un nom : remplacer
# « Dupont » partout dans un devis parce que c'est le client d'origine
# détruirait « rue Dupont » et « société Dupont & Fils ».
_VARIABLE = re.compile(r"\{\{?\s*([A-Za-zÀ-ÿ0-9_ .-]{1,40})\s*\}?\}"
                       r"|\[\[\s*([A-Za-zÀ-ÿ0-9_ .-]{1,40})\s*\]\]")


def _variables(textes: list[str]) -> list[str]:
    """Les noms de variables trouvés, dans l'ordre d'apparition, sans doublon."""
    vues: list[str] = []
    for t in textes:
        for m in _VARIABLE.finditer(t or ""):
            nom = (m.group(1) or m.group(2) or "").strip()
            if nom and nom not in vues:
                vues.append(nom)
    return vues


def _textes_docx(doc) -> list[str]:
    """Tous les textes d'un document Word, EN-TÊTES ET PIEDS COMPRIS.

    Les oublier serait passer à côté de ce qui porte le plus souvent l'identité
    du document : le logo, la raison sociale, le numéro de TVA, la pagination.
    """
    out: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for t in doc.tables:
        for ligne in t.rows:
            for cellule in ligne.cells:
                if cellule.text.strip():
                    out.append(cellule.text)
    for section in doc.sections:
        for zone in (section.header, section.footer):
            if zone is None:
                continue
            for p in zone.paragraphs:
                if p.text.strip():
                    out.append(p.text)
    return out


def analyser(octets: bytes, genre: str) -> dict:
    """Ce que porte le document : sa structure, son texte, ses variables.

    Sert à MONTRER une trame avant de l'enregistrer, et à dire ce qu'on saura
    y remplacer. Ne modifie rien.
    """
    if genre == "docx":
        import docx  # python-docx, déjà dans requirements

        doc = docx.Document(io.BytesIO(octets))
        textes = _textes_docx(doc)
        images = len(doc.inline_shapes)
        # `sections` porte les en-têtes ; on dit s'il y en a, parce que c'est
        # ce qui distingue un papier à en-tête d'une page blanche.
        entete = any((s.header is not None
                      and any(p.text.strip() for p in s.header.paragraphs))
                     for s in doc.sections)
        return {
            "genre": "docx",
            "paragraphes": len(doc.paragraphs),
            "tableaux": len(doc.tables),
            "images": images,
            "entete": entete,
            "textes": textes,
            "variables": _variables(textes),
        }

    if genre == "xlsx":
        from openpyxl import load_workbook

        classeur = load_workbook(io.BytesIO(octets), data_only=False)
        textes: list[str] = []
        feuilles = []
        for feuille in classeur.worksheets:
            lignes = feuille.max_row or 0
            colonnes = feuille.max_column or 0
            feuilles.append({"nom": feuille.title, "lignes": lignes,
                             "colonnes": colonnes})
            for ligne in feuille.iter_rows():
                for cellule in ligne:
                    v = cellule.value
                    if isinstance(v, str) and v.strip():
                        textes.append(v)
        return {
            "genre": "xlsx",
            "feuilles": feuilles,
            "textes": textes,
            "variables": _variables(textes),
        }

    raise ValueError(f"Type de trame non géré : {genre!r}")


def exploitable(analyse: dict) -> tuple[bool, str]:
    """Cette trame pourra-t-elle servir ? Et sinon, pourquoi ?"""
    texte = "".join(analyse.get("textes") or [])
    if len(texte.strip()) < MIN_TEXTE_UTILE:
        return False, ("Ce document ne contient presque aucun texte : c'est "
                       "probablement un scan ou une image. On ne pourra rien y "
                       "remplacer.")
    return True, ""


# ── Remplir : l'original, avec un autre contenu ──────────────────────────

def _remplacer_dans_paragraphe(paragraphe, table: dict) -> int:
    """Remplace dans UN paragraphe Word, en préservant sa mise en forme.

    LE PIÈGE DES RUNS, et la raison d'être de cette fonction. Word découpe un
    paragraphe en fragments (« runs ») à chaque changement de mise en forme, et
    en crée aussi pour ses propres raisons : une correction, un copier-coller,
    une reprise de frappe. « Devis n° DEV-2025-014 » vit donc souvent en
    plusieurs runs, et chercher la référence dans chacun ne trouve RIEN. C'est
    exactement sur les documents retouchés à la main — c'est-à-dire les vrais —
    qu'un remplacement naïf échoue en silence.

    On travaille donc sur le texte ENTIER du paragraphe. Si rien ne change, on
    ne touche à rien : ne pas réécrire un paragraphe intact, c'est garantir que
    sa mise en forme survit exactement.
    """
    avant = paragraphe.text
    apres = avant
    for cherche, remplace in table.items():
        if cherche and cherche in apres:
            apres = apres.replace(cherche, remplace)
    if apres == avant:
        return 0
    runs = paragraphe.runs
    if not runs:
        # Un paragraphe sans run (rare, mais existe) : on écrit directement.
        paragraphe.text = apres
        return 1
    # Le premier run garde SA mise en forme et reçoit tout le texte ; les
    # suivants sont vidés sans être supprimés (retirer un run d'un paragraphe
    # Word peut emporter avec lui des propriétés de la ligne).
    runs[0].text = apres
    for run in runs[1:]:
        run.text = ""
    return 1


def remplir(octets: bytes, genre: str, table: dict) -> tuple[bytes, int]:
    """L'original, avec les textes de `table` remplacés. Rend (octets, nombre).

    Tout ce qui n'est pas dans la table reste STRICTEMENT intact : styles,
    en-têtes, pieds de page, images, largeurs de colonnes, formules Excel. On
    n'a rien reconstruit, on a modifié.
    """
    if not isinstance(table, dict) or not table:
        raise ValueError("Aucun remplacement demandé.")
    if len(table) > MAX_REMPLACEMENTS:
        raise ValueError(f"Trop de remplacements ({len(table)}, plafond "
                         f"{MAX_REMPLACEMENTS}).")
    # Les clés vides remplaceraient PARTOUT : on les écarte avant d'ouvrir quoi
    # que ce soit, plutôt que de découvrir le dégât dans le fichier rendu.
    propre = {str(k): str(v) for k, v in table.items() if str(k).strip()}
    if not propre:
        raise ValueError("Aucun texte à chercher n'a été fourni.")

    faits = 0
    sortie = io.BytesIO()

    if genre == "docx":
        import docx

        doc = docx.Document(io.BytesIO(octets))
        for p in doc.paragraphs:
            faits += _remplacer_dans_paragraphe(p, propre)
        for t in doc.tables:
            for ligne in t.rows:
                for cellule in ligne.cells:
                    for p in cellule.paragraphs:
                        faits += _remplacer_dans_paragraphe(p, propre)
        # EN-TÊTES ET PIEDS AUSSI : c'est là que vivent la date, la référence
        # et le nom du client sur la plupart des documents d'entreprise. Les
        # oublier produirait un document qui se contredit lui-même, l'ancienne
        # référence subsistant en haut de chaque page.
        for section in doc.sections:
            for zone in (section.header, section.footer):
                if zone is None:
                    continue
                for p in zone.paragraphs:
                    faits += _remplacer_dans_paragraphe(p, propre)
                for t in zone.tables:
                    for ligne in t.rows:
                        for cellule in ligne.cells:
                            for p in cellule.paragraphs:
                                faits += _remplacer_dans_paragraphe(p, propre)
        doc.save(sortie)
        return sortie.getvalue(), faits

    if genre == "xlsx":
        from openpyxl import load_workbook

        # `data_only=False` GARDE LES FORMULES. À True, openpyxl ne lirait que
        # la dernière valeur calculée par Excel et les écrirait en dur : le
        # classeur rendu serait mort, ses totaux figés.
        classeur = load_workbook(io.BytesIO(octets), data_only=False)
        for feuille in classeur.worksheets:
            for ligne in feuille.iter_rows():
                for cellule in ligne:
                    v = cellule.value
                    if not isinstance(v, str):
                        continue
                    neuf = v
                    for cherche, remplace in propre.items():
                        if cherche in neuf:
                            neuf = neuf.replace(cherche, remplace)
                    if neuf != v:
                        # Écrire la cellule ne touche pas à son style :
                        # openpyxl les porte séparément de la valeur.
                        cellule.value = neuf
                        faits += 1
        classeur.save(sortie)
        return sortie.getvalue(), faits

    raise ValueError(f"Type de trame non géré : {genre!r}")
