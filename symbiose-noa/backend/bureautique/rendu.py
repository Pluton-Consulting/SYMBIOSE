"""
Rendu d'un document : même description, trois formats.

Le flux d'éléments est parcouru UNE fois et écrit au fil de l'eau. Rien n'est
accumulé : c'est ce qui permet un document de plusieurs centaines de pages sans
rapport avec la mémoire disponible.

Chaque format honore ce qu'il sait faire :
  * .docx  en-tête et pied de section, numérotation par champ Word, styles de
           titre natifs (donc sommaire automatique possible côté Word).
  * .pdf   en-tête et pied dessinés sur CHAQUE page, numérotation calculée au
           tirage, tableaux qui se poursuivent d'une page à l'autre.
  * .xlsx  un onglet par « feuille », en-tête et pied d'impression, en-têtes de
           colonnes figées.

Un bloc qu'un format ne peut pas honorer est DÉGRADÉ, jamais ignoré : une
« feuille » dans un .docx devient un tableau précédé de son nom. Le lecteur voit
alors une mise en forme imparfaite plutôt qu'un trou, et un trou est ce qu'on ne
remarque qu'une fois le document envoyé.
"""
from __future__ import annotations

import logging

logger = logging.getLogger("symbiose.bureautique.rendu")


def rendre(entete: dict, elements, sortie: str) -> str:
    fmt = entete.get("format", "docx")
    if fmt == "xlsx":
        return _xlsx(entete, elements, sortie)
    if fmt == "pdf":
        return _pdf(entete, elements, sortie)
    return _docx(entete, elements, sortie)


# ── Word ──────────────────────────────────────────────────────────────
def _docx(entete: dict, elements, sortie: str) -> str:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm, RGBColor

    from bureautique.modele import COULEURS, TAILLES

    doc = Document()
    section = doc.sections[0]
    if entete.get("paysage"):
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    for marge in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, marge, Cm(2))

    if entete.get("entete"):
        p = section.header.paragraphs[0]
        p.text = entete["entete"]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if entete.get("pied") or entete.get("numeroter"):
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if entete.get("pied"):
            p.add_run(entete["pied"] + ("   ·   " if entete.get("numeroter") else ""))
        if entete.get("numeroter"):
            # Champ Word : Word calcule la pagination à l'ouverture. Écrire un
            # numéro en dur donnerait « page 1 » sur toutes les pages.
            _champ_page(p)

    doc.add_heading(entete["titre"], 0)
    if entete.get("sous_titre"):
        p = doc.add_paragraph(entete["sous_titre"])
        p.runs[0].font.size = Pt(13)
        p.runs[0].font.italic = True

    for e in elements:
        bloc = e["bloc"]
        if bloc == "titre":
            doc.add_heading(e["texte"], e["niveau"])
        elif bloc == "paragraphe":
            p = doc.add_paragraph()
            run = p.add_run(e["texte"])
            run.font.bold = e.get("gras", False)
            run.font.italic = e.get("italique", False)
            run.font.size = Pt(TAILLES.get(e.get("taille"), 11))
            if e.get("couleur"):
                run.font.color.rgb = RGBColor.from_string(COULEURS[e["couleur"]])
            if e.get("centre"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif bloc == "liste":
            style = "List Number" if e["ordonnee"] else "List Bullet"
            for item in e["items"]:
                doc.add_paragraph(item, style=style)
        elif bloc in ("tableau", "feuille"):
            if bloc == "feuille":
                doc.add_heading(e.get("nom") or "Feuille", 2)
            _tableau_docx(doc, e)
            if e.get("legende"):
                p = doc.add_paragraph(e["legende"])
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.italic = True
        elif bloc == "saut_page":
            doc.add_page_break()
        elif bloc == "separateur":
            doc.add_paragraph("―" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(sortie)
    return sortie


def _champ_page(paragraphe) -> None:
    """Insère « page X sur Y » sous forme de champs Word."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def champ(code: str):
        debut = OxmlElement("w:fldChar"); debut.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
        instr.text = code
        fin = OxmlElement("w:fldChar"); fin.set(qn("w:fldCharType"), "end")
        run = paragraphe.add_run()
        for e in (debut, instr, fin):
            run._r.append(e)

    champ(" PAGE ")
    paragraphe.add_run(" / ")
    champ(" NUMPAGES ")


def _tableau_docx(doc, e: dict) -> None:
    entetes, lignes = e["entetes"], e["lignes"]
    colonnes = len(entetes) or (len(lignes[0]) if lignes else 0)
    if not colonnes:
        return
    table = doc.add_table(rows=1 if entetes else 0, cols=colonnes)
    table.style = "Light Grid Accent 1"
    if entetes:
        for i, titre in enumerate(entetes):
            cellule = table.rows[0].cells[i]
            cellule.text = titre
            for p in cellule.paragraphs:
                for r in p.runs:
                    r.font.bold = True
    for ligne in lignes:
        cellules = table.add_row().cells
        for i, valeur in enumerate(ligne[:colonnes]):
            cellules[i].text = valeur


# ── PDF ───────────────────────────────────────────────────────────────
def _pdf(entete: dict, elements, sortie: str) -> str:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, PageBreak, HRFlowable)

    format_page = landscape(A4) if entete.get("paysage") else A4
    styles = getSampleStyleSheet()
    petit = ParagraphStyle("petit", parent=styles["Normal"], fontSize=8,
                           textColor=colors.grey)
    legende = ParagraphStyle("legende", parent=petit, alignment=TA_CENTER)

    def decor(canvas, doc):
        """En-tête et pied DESSINÉS sur chaque page, numérotation au tirage."""
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.grey)
        largeur, hauteur = format_page
        if entete.get("entete"):
            canvas.drawRightString(largeur - 2 * cm, hauteur - 1.2 * cm, entete["entete"])
        bas = []
        if entete.get("pied"):
            bas.append(entete["pied"])
        if entete.get("numeroter"):
            bas.append(f"page {canvas.getPageNumber()}")
        if bas:
            canvas.drawCentredString(largeur / 2, 1.2 * cm, "   ·   ".join(bas))
        canvas.restoreState()

    doc = SimpleDocTemplate(sortie, pagesize=format_page,
                            topMargin=2.2 * cm, bottomMargin=2.2 * cm,
                            leftMargin=2 * cm, rightMargin=2 * cm,
                            title=entete["titre"])

    flux = [Paragraph(entete["titre"], styles["Title"])]
    if entete.get("sous_titre"):
        flux.append(Paragraph(entete["sous_titre"], styles["Italic"]))
    flux.append(Spacer(1, 0.6 * cm))

    for e in elements:
        bloc = e["bloc"]
        if bloc == "titre":
            flux.append(Spacer(1, 0.35 * cm))
            flux.append(Paragraph(_echapper(e["texte"]),
                                  styles[f"Heading{min(e['niveau'], 4)}"]))
        elif bloc == "paragraphe":
            flux.append(Paragraph(_echapper(e["texte"]), _style_paragraphe(e, styles)))
        elif bloc == "liste":
            for i, item in enumerate(e["items"], 1):
                puce = f"{i}." if e["ordonnee"] else "•"
                flux.append(Paragraph(f"{puce}&nbsp;&nbsp;{_echapper(item)}",
                                      styles["BodyText"]))
        elif bloc in ("tableau", "feuille"):
            if bloc == "feuille":
                flux.append(Paragraph(_echapper(e.get("nom") or "Feuille"),
                                      styles["Heading2"]))
            table = _tableau_pdf(e, styles, format_page)
            if table is not None:
                flux.append(table)
            if e.get("legende"):
                flux.append(Paragraph(_echapper(e["legende"]), legende))
            flux.append(Spacer(1, 0.35 * cm))
        elif bloc == "saut_page":
            flux.append(PageBreak())
        elif bloc == "separateur":
            flux.append(Spacer(1, 0.25 * cm))
            flux.append(HRFlowable(width="100%", color=colors.lightgrey))
            flux.append(Spacer(1, 0.25 * cm))

    doc.build(flux, onFirstPage=decor, onLaterPages=decor)
    return sortie


def _style_paragraphe(e: dict, styles):
    """Style PDF d'un paragraphe, d'après sa mise en forme demandée.

    Un style est construit à la volée plutôt que puisé dans une table : les
    combinaisons (gras × taille × couleur × centré) sont trop nombreuses pour
    être toutes prévues, et un style manquant rendrait le texte silencieusement
    sans sa mise en forme.
    """
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import ParagraphStyle

    from bureautique.modele import COULEURS, TAILLES

    taille = TAILLES.get(e.get("taille"), 11)
    style = ParagraphStyle(
        f"p{taille}{e.get('couleur','')}{e.get('gras')}{e.get('centre')}",
        parent=styles["BodyText"],
        fontSize=taille,
        # L'interligne doit suivre la taille, sinon un texte en 36 points se
        # chevauche d'une ligne à l'autre.
        leading=taille * 1.25,
        alignment=TA_CENTER if e.get("centre") else styles["BodyText"].alignment,
    )
    if e.get("couleur"):
        style.textColor = colors.HexColor("#" + COULEURS[e["couleur"]])
    gras, italique = e.get("gras"), e.get("italique")
    if gras and italique:
        style.fontName = "Helvetica-BoldOblique"
    elif gras:
        style.fontName = "Helvetica-Bold"
    elif italique:
        style.fontName = "Helvetica-Oblique"
    return style


def _echapper(texte: str) -> str:
    """Le texte va dans un moteur qui interprète des balises : on le neutralise."""
    return (texte.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _tableau_pdf(e: dict, styles, format_page):
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, Table, TableStyle

    entetes, lignes = e["entetes"], e["lignes"]
    if not entetes and not lignes:
        return None

    cellule = ParagraphStyle("cellule", parent=styles["BodyText"], fontSize=8, leading=10)
    entete_style = ParagraphStyle("entete", parent=cellule, textColor=colors.white)

    donnees = []
    if entetes:
        donnees.append([Paragraph(f"<b>{_echapper(c)}</b>", entete_style) for c in entetes])
    colonnes = len(entetes) or max((len(l) for l in lignes), default=0)
    for ligne in lignes:
        ligne = list(ligne) + [""] * (colonnes - len(ligne))
        donnees.append([Paragraph(_echapper(c), cellule) for c in ligne[:colonnes]])
    if not donnees:
        return None

    # Largeur répartie sur la place utile : sans cela, un tableau large déborde
    # de la page et les dernières colonnes sont invisibles à l'impression.
    utile = format_page[0] - 4 * cm
    table = Table(donnees, colWidths=[utile / colonnes] * colonnes,
                  repeatRows=1 if entetes else 0)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5233") if entetes else colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B0B0B0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F6F4")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


# ── Excel ─────────────────────────────────────────────────────────────
def _xlsx(entete: dict, elements, sortie: str) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    from bureautique.modele import MAX_FEUILLES, COULEURS, TAILLES

    classeur = Workbook()
    classeur.remove(classeur.active)
    gras_blanc = Font(bold=True, color="FFFFFF")
    fond = PatternFill("solid", fgColor="2F5233")

    feuille = None
    ligne_courante = 1
    noms = set()

    def nouvelle(nom: str):
        nonlocal feuille, ligne_courante
        if len(classeur.worksheets) >= MAX_FEUILLES:
            return feuille
        # Excel refuse les doublons et certains caractères dans un nom d'onglet.
        propre = "".join(c for c in (nom or "Feuille") if c not in "[]:*?/\\")[:31] or "Feuille"
        base, n = propre, 2
        while propre in noms:
            propre = f"{base[:28]}_{n}"
            n += 1
        noms.add(propre)
        feuille = classeur.create_sheet(propre)
        if entete.get("entete"):
            feuille.oddHeader.right.text = entete["entete"]
        if entete.get("pied") or entete.get("numeroter"):
            bas = [entete.get("pied") or ""]
            if entete.get("numeroter"):
                bas.append("page &P / &N")
            feuille.oddFooter.center.text = "   ·   ".join(x for x in bas if x)
        ligne_courante = 1
        return feuille

    def ecrire(valeurs, gras=False):
        nonlocal ligne_courante
        for i, v in enumerate(valeurs, 1):
            c = feuille.cell(row=ligne_courante, column=i, value=v)
            if gras:
                c.font, c.fill = gras_blanc, fond
            c.alignment = Alignment(vertical="top", wrap_text=True)
        ligne_courante += 1

    def garantir():
        """Ouvre la feuille d'accueil, mais SEULEMENT si quelque chose y va.

        AVANT (relevé de Noa, 01/09) : « à chaque fois que je demande des Excel,
        il me fait toujours deux feuilles — une avec un titre qui est inutile et
        une autre avec les vraies infos ». La feuille d'accueil était créée
        d'office, puis chaque bloc `feuille` créait la sienne : un classeur avec
        un onglet ne portant qu'un titre, et un onglet de données à côté.

        Elle n'est donc plus ouverte qu'à la demande. Un document fait
        uniquement de blocs `feuille` — le cas de tous les exports de listes —
        n'en a plus du tout, et l'onglet porte directement le nom des données.
        """
        nonlocal ligne_courante
        if feuille is not None:
            return
        nouvelle(entete["titre"])
        ecrire([entete["titre"]], gras=False)
        feuille.cell(row=1, column=1).font = Font(bold=True, size=14)
        ligne_courante = 3

    for e in elements:
        bloc = e["bloc"]
        if bloc != "feuille":
            garantir()
        if bloc == "feuille":
            nouvelle(e.get("nom") or "Feuille")
            if e["entetes"]:
                ecrire(e["entetes"], gras=True)
                feuille.freeze_panes = "A2"     # les entêtes restent visibles
            for ligne in e["lignes"]:
                ecrire(ligne)
        elif bloc == "tableau":
            if e.get("legende"):
                ecrire([e["legende"]])
            if e["entetes"]:
                ecrire(e["entetes"], gras=True)
            for ligne in e["lignes"]:
                ecrire(ligne)
            ligne_courante += 1
        elif bloc == "titre":
            ecrire([e["texte"]])
            feuille.cell(row=ligne_courante - 1, column=1).font = Font(
                bold=True, size=max(14 - e["niveau"], 10))
        elif bloc == "paragraphe":
            ecrire([e["texte"]])
            c = feuille.cell(row=ligne_courante - 1, column=1)
            c.font = Font(bold=e.get("gras", False), italic=e.get("italique", False),
                          size=TAILLES.get(e.get("taille"), 11),
                          color=(COULEURS[e["couleur"]] if e.get("couleur") else None))
            if e.get("centre"):
                c.alignment = Alignment(horizontal="center", vertical="top", wrap_text=True)
        elif bloc == "liste":
            for i, item in enumerate(e["items"], 1):
                ecrire([f"{i}." if e["ordonnee"] else "•", item])
        elif bloc in ("saut_page", "separateur"):
            ligne_courante += 1

    # Un classeur SANS AUCUNE feuille est un fichier qu'Excel refuse d'ouvrir :
    # si le document était vide, on ouvre la feuille d'accueil pour de bon.
    garantir()

    # Largeurs : un classeur dont tout est tronqué à l'écran passe pour cassé.
    for ws in classeur.worksheets:
        for colonne in range(1, min(ws.max_column, 40) + 1):
            largeur = max((len(str(ws.cell(row=r, column=colonne).value or ""))
                           for r in range(1, min(ws.max_row, 200) + 1)), default=10)
            ws.column_dimensions[get_column_letter(colonne)].width = min(max(largeur + 2, 10), 60)

    classeur.save(sortie)
    return sortie
