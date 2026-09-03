"""
Lecture des fichiers déposés à l'import : CSV, Excel, Word, PDF, texte.

Deux familles de résultats, parce qu'elles ne s'ingèrent pas pareil :

  * TABULAIRE (csv, xlsx, xls) -> une LIGNE = un document. Un export de milliers
    de lignes ingéré en un seul bloc serait découpé arbitrairement et la recherche
    remonterait des morceaux sans rapport ; ligne par ligne, chaque chantier /
    devis / client devient un document retrouvable, avec un identifiant stable.

  * DOCUMENT (pdf, docx, txt, md) -> un FICHIER = un document, découpé ensuite
    par le pipeline d'ingestion.

Toutes les fonctions sont synchrones et bornées (nb de lignes, nb de pages) :
les appeler via asyncio.to_thread pour ne pas bloquer la boucle d'événements.
"""
from __future__ import annotations

import csv
import io
import logging
from typing import Optional

logger = logging.getLogger("symbiose.ingestion.parsers")

# Bornes anti-fichier piégé / anti-saturation mémoire.
MAX_LIGNES = 20000
MAX_PAGES_PDF = 300
MAX_PAGES_OCR = 60        # l'OCR coûte ~1 s/page : borne plus basse que la lecture texte
MAX_COLONNES = 200

# En dessous de ce nombre de caractères par page, on considère que le PDF n'a pas
# de couche texte (document scanné) et on bascule sur l'OCR.
SEUIL_TEXTE_PAR_PAGE = 40

OCR_DPI = 200             # compromis lisibilité / mémoire pour le rendu des pages
OCR_LANGUES = "fra+eng"   # documents FR, mais les CCTP contiennent souvent de l'anglais

ENCODAGES = ("utf-8-sig", "utf-8", "cp1252", "latin-1")

EXT_TABULAIRE = (".csv", ".xlsx", ".xls", ".xlsm")
EXT_DOCUMENT = (".pdf", ".docx", ".txt", ".md")
EXT_IMAGE = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp")


class FichierNonSupporte(Exception):
    """Extension inconnue ou dépendance de lecture absente."""


def famille(nom: str) -> Optional[str]:
    """'tabulaire', 'document', ou None si l'extension n'est pas gérée."""
    n = (nom or "").lower()
    if n.endswith(EXT_TABULAIRE):
        return "tabulaire"
    if n.endswith(EXT_DOCUMENT) or n.endswith(EXT_IMAGE):
        return "document"
    return None


# ── OCR (documents scannés, photos de documents) ────────────────────────────
# Volontairement LOCAL (tesseract) : une facture ou un CCTP scanné ne doit pas
# être envoyé à un service de reconnaissance externe. Absence de tesseract =
# dégradation propre, jamais de plantage.

def ocr_disponible() -> bool:
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def ocr_image(brut: bytes) -> str:
    """Texte d'une image (photo ou scan d'un document)."""
    try:
        import io as _io
        import pytesseract
        from PIL import Image
    except ImportError as e:
        raise FichierNonSupporte("OCR indisponible (pytesseract/Pillow absent)") from e
    try:
        with Image.open(_io.BytesIO(brut)) as img:
            if img.mode not in ("L", "RGB"):        # CMJN / palette / alpha -> RGB
                img = img.convert("RGB")
            return (pytesseract.image_to_string(img, lang=OCR_LANGUES) or "").strip()
    except Exception as e:
        raise FichierNonSupporte(f"OCR de l'image impossible : {e}") from e


def ocr_pdf(brut: bytes) -> str:
    """Rend chaque page en image puis l'OCRise. Utilisé quand le PDF n'a pas de
    couche texte (document scanné)."""
    try:
        import pypdfium2 as pdfium
        import pytesseract
    except ImportError as e:
        raise FichierNonSupporte(
            "OCR PDF indisponible (pypdfium2/pytesseract absent)"
        ) from e

    morceaux = []
    doc = pdfium.PdfDocument(brut)
    try:
        total = min(len(doc), MAX_PAGES_OCR)
        if len(doc) > MAX_PAGES_OCR:
            logger.warning("OCR limité aux %d premières pages (sur %d)", MAX_PAGES_OCR, len(doc))
        for i in range(total):
            page = doc[i]
            image = page.render(scale=OCR_DPI / 72).to_pil()
            try:
                morceaux.append(pytesseract.image_to_string(image, lang=OCR_LANGUES) or "")
            finally:
                image.close()
    finally:
        doc.close()
    return "\n\n".join(morceaux).strip()


def _decoder(brut: bytes) -> str:
    """Décode en essayant les encodages usuels (Excel FR écrit en cp1252)."""
    for enc in ENCODAGES:
        try:
            return brut.decode(enc)
        except UnicodeDecodeError:
            continue
    return brut.decode("utf-8", errors="replace")


def lire_csv(brut: bytes) -> tuple[list[str], list[dict]]:
    texte = _decoder(brut)
    echantillon = texte[:8192]
    try:
        sep = csv.Sniffer().sniff(echantillon, delimiters=";,\t|").delimiter
    except csv.Error:
        sep = ";" if echantillon.count(";") > echantillon.count(",") else ","

    lecteur = csv.DictReader(texte.splitlines(), delimiter=sep)
    entetes = [(c or "").strip() for c in (lecteur.fieldnames or [])][:MAX_COLONNES]
    lignes = []
    for i, l in enumerate(lecteur):
        if i >= MAX_LIGNES:
            logger.warning("CSV tronqué à %d lignes", MAX_LIGNES)
            break
        if any((v or "").strip() for v in l.values()):
            lignes.append({(k or "").strip(): (v or "").strip() for k, v in l.items() if k})
    return entetes, lignes


def format_tabulaire(brut: bytes) -> str:
    """Format RÉEL d'un fichier tabulaire, d'après son contenu.

    L'extension ment souvent. Les logiciels de gestion exportent couramment un
    tableau HTML ou un CSV sous un nom en `.xls`, et un `.xls` authentique n'est
    pas du tout un `.xlsx` : le premier est un conteneur OLE2, le second une
    archive zip. Se fier au nom faisait passer tout cela à openpyxl, qui
    répondait « File is not a zip file » — exact, mais incompréhensible pour qui
    vient d'exporter depuis son logiciel métier.
    """
    tete = brut[:8]
    if tete[:4] == b"PK\x03\x04":
        return "xlsx"                      # archive zip : OOXML
    if tete[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        return "xls"                       # conteneur OLE2 : Excel 97-2003
    debut = brut[:4096].lstrip()[:512].lower()
    if debut.startswith(b"<") and (b"<table" in brut[:65536].lower()
                                   or b"<html" in debut or b"<?xml" in debut):
        return "html"
    return "csv"                           # texte délimité, quel que soit le nom


def lire_html(brut: bytes) -> tuple[list[str], list[dict]]:
    """Lit le premier tableau d'un export HTML déguisé en tableur.

    Écrit sur la bibliothèque standard : `pandas.read_html` exigerait lxml,
    bs4 ou html5lib, absents ici. Un tableau d'export est plat — pas de
    tableaux imbriqués, pas de mise en forme — donc un analyseur simple suffit
    et évite d'ajouter trois dépendances pour un cas de compatibilité.
    """
    from html.parser import HTMLParser

    class _Tableau(HTMLParser):
        def __init__(self):
            super().__init__(convert_charrefs=True)
            self.lignes: list[list[str]] = []
            self._ligne: list[str] | None = None
            self._cellule: list[str] | None = None

        def handle_starttag(self, tag, attrs):
            if tag == "tr":
                self._ligne = []
            elif tag in ("td", "th") and self._ligne is not None:
                self._cellule = []
            elif tag == "br" and self._cellule is not None:
                self._cellule.append(" ")

        def handle_endtag(self, tag):
            if tag in ("td", "th") and self._cellule is not None:
                self._ligne.append(" ".join("".join(self._cellule).split()))
                self._cellule = None
            elif tag == "tr" and self._ligne is not None:
                if any(c for c in self._ligne):
                    self.lignes.append(self._ligne[:MAX_COLONNES])
                self._ligne = None

        def handle_data(self, data):
            if self._cellule is not None:
                self._cellule.append(data)

    analyseur = _Tableau()
    # Pas de `unescape` avant l'analyse : `convert_charrefs` décode déjà les
    # entités DANS le texte des cellules. Décoder en amont transformerait un
    # « &lt;table&gt; » écrit dans une cellule en vraie balise, et disloquerait
    # la structure du tableau.
    analyseur.feed(_decoder(brut))
    if not analyseur.lignes:
        raise FichierNonSupporte(
            "Ce fichier ressemble à une page HTML mais ne contient aucun tableau. "
            "Réexportez-le en CSV ou en Excel (.xlsx).")

    entetes = [c.strip() for c in analyseur.lignes[0]]
    lignes = []
    for ligne in analyseur.lignes[1:MAX_LIGNES + 1]:
        d = {e: v.strip() for e, v in zip(entetes, ligne) if e and v.strip()}
        if d:
            lignes.append(d)
    return [e for e in entetes if e], lignes


def lire_xls(brut: bytes) -> tuple[list[str], list[dict]]:
    """Excel 97-2003 (.xls authentique). openpyxl ne sait pas le lire."""
    try:
        import xlrd
    except ImportError as e:
        raise FichierNonSupporte(
            "Ce fichier est un Excel ancien (.xls, format 97-2003) que le serveur "
            "ne sait pas lire. Ouvrez-le puis « Enregistrer sous » en .xlsx ou en "
            "CSV, et réimportez-le.") from e

    classeur = xlrd.open_workbook(file_contents=brut)
    feuille = classeur.sheet_by_index(0)
    if not feuille.nrows:
        raise FichierNonSupporte("Fichier Excel vide (aucune ligne).")

    def _texte(c) -> str:
        return "" if c is None else " ".join(str(c).split())

    entetes = [_texte(v) for v in feuille.row_values(0)][:MAX_COLONNES]
    lignes = []
    for i in range(1, min(feuille.nrows, MAX_LIGNES + 1)):
        d = {e: _texte(v) for e, v in zip(entetes, feuille.row_values(i))
             if e and _texte(v)}
        if d:
            lignes.append(d)
    return [e for e in entetes if e], lignes


def lire_excel(brut: bytes) -> tuple[list[str], list[dict]]:
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise FichierNonSupporte("Lecture Excel indisponible (openpyxl absent)") from e

    wb = load_workbook(io.BytesIO(brut), read_only=True, data_only=True)
    ws = wb.active                      # 1re feuille — la plus courante pour un export
    iterateur = ws.iter_rows(values_only=True)

    entetes: list[str] = []
    for ligne in iterateur:             # 1re ligne non vide = en-têtes
        if ligne and any(c is not None and str(c).strip() for c in ligne):
            entetes = [str(c).strip() if c is not None else "" for c in ligne][:MAX_COLONNES]
            break

    # UNE COLONNE SANS EN-TÊTE N'EST PAS UNE COLONNE VIDE (03/09). Le fichier
    # « tableau client entretien.xlsx » de Symbiose : 8 colonnes nommées à
    # gauche, puis — pour la moitié des lignes, collées depuis un autre export —
    # l'adresse mail en colonne AG, sous un en-tête VIDE. Chaque cellule sans
    # en-tête était jetée : ces clients existaient sans adresse, et le
    # publipostage n'en voyait que la moitié. Une colonne qui porte une valeur
    # reçoit le nom de sa lettre Excel ; le modèle et les actions la voient.
    vues: set[str] = set()
    for i, e in enumerate(entetes):
        if e and e in vues:             # deux colonnes « Nom » : la seconde se distingue
            entetes[i] = f"{e} ({_lettre_colonne(i)})"
        vues.add(entetes[i])

    lignes = []
    for i, ligne in enumerate(iterateur):
        if i >= MAX_LIGNES:
            logger.warning("Excel tronqué à %d lignes", MAX_LIGNES)
            break
        if not ligne or not any(c is not None and str(c).strip() for c in ligne):
            continue
        d = {}
        for j, valeur in enumerate(ligne[:MAX_COLONNES]):
            if valeur is None or not str(valeur).strip():
                continue
            entete = entetes[j] if j < len(entetes) else ""
            if not entete:
                entete = f"Colonne {_lettre_colonne(j)}"
                if entete not in entetes:
                    entetes.append(entete)
            d[entete] = str(valeur).strip()
        if d:
            lignes.append(d)
    wb.close()
    return [e for e in entetes if e], lignes


def _lettre_colonne(indice: int) -> str:
    """0 → A, 25 → Z, 26 → AA : la lettre qu'Excel montre en tête de colonne."""
    lettres = ""
    n = indice
    while True:
        lettres = chr(ord("A") + n % 26) + lettres
        n = n // 26 - 1
        if n < 0:
            return lettres


def lire_pdf(brut: bytes) -> str:
    """Texte d'un PDF. Bascule automatiquement sur l'OCR si le fichier n'a pas de
    couche texte exploitable (PDF scanné : chaque page n'est qu'une image)."""
    try:
        import pdfplumber
    except ImportError as e:
        raise FichierNonSupporte("Lecture PDF indisponible (pdfplumber absent)") from e

    morceaux = []
    with pdfplumber.open(io.BytesIO(brut)) as pdf:
        pages_lues = min(len(pdf.pages), MAX_PAGES_PDF)
        if len(pdf.pages) > MAX_PAGES_PDF:
            logger.warning("PDF tronqué à %d pages", MAX_PAGES_PDF)
        for page in pdf.pages[:pages_lues]:
            morceaux.append(page.extract_text() or "")

    texte = "\n\n".join(morceaux).strip()
    if pages_lues and len(texte) >= SEUIL_TEXTE_PAR_PAGE * pages_lues:
        return texte

    # Trop peu de texte pour le nombre de pages -> probablement scanné.
    # Règle de sûreté : l'OCR ne doit JAMAIS faire perdre du texte déjà extrait.
    # Un PDF court mais authentiquement textuel (note d'une ligne, courrier bref)
    # tombe sous le seuil : on le renvoie tel quel plutôt que d'échouer.
    if not ocr_disponible():
        if texte:
            logger.info("PDF peu fourni (%d caractères) et OCR indisponible — texte brut conservé", len(texte))
            return texte
        raise FichierNonSupporte(
            "Ce PDF ne contient aucun texte (document scanné) et la reconnaissance "
            "de caractères n'est pas disponible sur le serveur."
        )

    logger.info("PDF sans couche texte (%d caractères / %d pages) — OCR", len(texte), pages_lues)
    try:
        ocr = ocr_pdf(brut)
    except FichierNonSupporte:
        if texte:
            return texte        # l'OCR a échoué : on garde ce qu'on avait
        raise
    return ocr if len(ocr) > len(texte) else texte


def lire_docx(brut: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise FichierNonSupporte("Lecture Word indisponible (python-docx absent)") from e
    d = docx.Document(io.BytesIO(brut))
    morceaux = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # Les tableaux Word portent souvent l'essentiel (métrés, postes, quantités).
    for table in d.tables:
        for row in table.rows:
            cellules = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cellules:
                morceaux.append(" | ".join(cellules))
    return "\n".join(morceaux).strip()


def ligne_en_texte(ligne: dict) -> str:
    """« Colonne : valeur » — lisible par le modèle, et les libellés participent
    eux-mêmes à la recherche sémantique."""
    return "\n".join(f"{k} : {v}" for k, v in ligne.items() if str(v).strip())


def analyser(nom: str, brut: bytes) -> dict:
    """Lit le fichier et retourne sa structure, sans rien écrire.

    Retour : {kind, columns, rows, text, documents_estimes}
    """
    fam = famille(nom)
    if fam is None:
        raise FichierNonSupporte(
            f"Format non géré : {nom}. Formats acceptés : CSV, Excel (.xlsx/.xls), "
            "Word (.docx), PDF (y compris scanné), images (.png/.jpg/.tif), texte (.txt/.md)."
        )

    n = nom.lower()
    if n.endswith(EXT_IMAGE):
        if not ocr_disponible():
            raise FichierNonSupporte(
                "Import d'image impossible : la reconnaissance de caractères "
                "n'est pas disponible sur le serveur."
            )
        texte = ocr_image(brut)
        if not texte:
            raise FichierNonSupporte(
                "Aucun texte reconnu dans cette image. Vérifiez qu'elle est nette, "
                "droite et bien éclairée."
            )
        return {"kind": "document", "columns": [], "rows": [],
                "text": texte, "documents_estimes": 1}

    if fam == "tabulaire":
        # Le CONTENU décide, pas l'extension : un export nommé `.xls` est très
        # souvent un tableau HTML ou un CSV, et un vrai `.xls` n'est pas un
        # `.xlsx`. Choisir d'après le nom envoyait tout à openpyxl, qui refusait
        # avec « File is not a zip file ».
        reel = format_tabulaire(brut)
        lecteurs = {"xlsx": lire_excel, "xls": lire_xls,
                    "html": lire_html, "csv": lire_csv}
        try:
            entetes, lignes = lecteurs[reel](brut)
        except FichierNonSupporte:
            raise
        except Exception as e:
            raise FichierNonSupporte(
                f"Lecture impossible : le fichier « {nom} » a été reconnu comme "
                f"{reel.upper()} d'après son contenu, mais n'a pas pu être lu ({e}). "
                "Réexportez-le en CSV ou en Excel (.xlsx).") from e
        if not lignes:
            raise FichierNonSupporte(
                f"Aucune ligne de données lisible dans « {nom} » (reconnu comme "
                f"{reel.upper()}). Vérifiez que la première ligne contient bien "
                "les en-têtes de colonnes.")
        return {"kind": "tabulaire", "columns": entetes, "rows": lignes,
                "text": None, "documents_estimes": len(lignes)}

    if n.endswith(".pdf"):
        texte = lire_pdf(brut)
    elif n.endswith(".docx"):
        texte = lire_docx(brut)
    else:
        texte = _decoder(brut).strip()

    if not texte:
        raise FichierNonSupporte("Aucun texte extractible de ce fichier.")
    return {"kind": "document", "columns": [], "rows": [],
            "text": texte, "documents_estimes": 1}
